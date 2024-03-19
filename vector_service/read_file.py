"""This module belong to all the service related to reading a file"""
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from docx import Document as Word_document
from fastapi import UploadFile
import pdfplumber
import traceback
import json
import io
import logging
import pandas as pd
logging.basicConfig(level=logging.INFO, format = '%(asctime)s - %(levelname)s -%(message)s')

MAX_CHUNK_SIZE = 100
CHUNK_OVERLAP = 10

async def read_pdf_file(file: UploadFile) -> list[Document]:
    """The function will read a file, chunk it and then will convert it to 'Document' Type Object and will return a list of 'Document' objectss.

    Args:
        file (UploadFile): the instance of uploaded file

    Returns:
        list[Document]: list of Document object after chunking
    """
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size = MAX_CHUNK_SIZE,
            chunk_overlap = CHUNK_OVERLAP
        )
        stream = io.BytesIO(await file.read())
        pdf = pdfplumber.open(stream)
        pages = pdf.pages
        final_list = []
        for idx, page in enumerate(pages):
            text_data = page.extract_text()
            for content in text_splitter.split_text(text_data):
                metadata = {
                    'file_name': file.filename,
                    'source_type':".pdf",
                    'page_number': idx
                }
                final_list.append(Document(page_content=content, metadata=metadata))
        return final_list
    except Exception as e:
       error_message = type(e).__name__
       traceback_message = traceback.format_exc()
       logging.error("Error: %s\nTraceback: %s", error_message, traceback_message)
       return []

async def read_text_file(file: UploadFile) -> list[Document]:
    """The function will read a file, chunk it and then will convert it to 'Document' Type Object and will return a list of 'Document' objectss.

    Args:
        file (UploadFile): the instance of uploaded file

    Returns:
        list[Document]: list of Document object after chunking
    """
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size = MAX_CHUNK_SIZE,
            chunk_overlap = CHUNK_OVERLAP
        )
        stream = await file.read()
        pdf_stream = io.BytesIO(stream)
        final_list = []
        data = pdf_stream.read().decode()
        for content in text_splitter.split_text(data):
            metadata = {
                'file_name': file.filename,
                'source_type':".txt",
                'content_length': len(content)
            }
            final_list.append(Document(page_content=content, metadata=metadata))
        return final_list
    except Exception as e:
       error_message = type(e).__name__
       traceback_message = traceback.format_exc()
       logging.error("Error: %s\nTraceback: %s", error_message, traceback_message)
       return []

async def read_csv_file(file: UploadFile) -> list[Document]:
    """The function will read a file, chunk it and then will convert it to 'Document' Type Object and will return a list of 'Document' objectss.

    Args:
        file (UploadFile): the instance of uploaded file

    Returns:
        list[Document]: list of Document object after chunking
    """
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size = MAX_CHUNK_SIZE,
            chunk_overlap = CHUNK_OVERLAP
        )
        df = pd.read_csv(file.file)
        rows = df.to_dict(orient='records')
        final_list = []
        for row in rows:
            data = json.dumps(row)
            for content in text_splitter.split_text(data):
                metadata = {
                    'file_name': file.filename,
                    'source_type':".csv"
                }
                final_list.append(Document(page_content=content, metadata=metadata))
        return final_list
    except Exception as e:
       error_message = type(e).__name__
       traceback_message = traceback.format_exc()
       logging.error("Error: %s\nTraceback: %s", error_message, traceback_message)
       return []
    
async def read_docx_file(file: UploadFile) -> list[Document]:
    """The function will read a file, chunk it and then will convert it to 'Document' Type Object and will return a list of 'Document' objectss.

    Args:
        file (UploadFile): the instance of uploaded file

    Returns:
        list[Document]: list of Document object after chunking
    """
    try:
        stream = await file.read()
        pdf_stream = io.BytesIO(stream)
        doc = Word_document(pdf_stream)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        all_content =  '\n'.join(text)
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size = MAX_CHUNK_SIZE,
            chunk_overlap = CHUNK_OVERLAP
        )
        final_list = []
        for content in text_splitter.split_text(all_content):
            metadata = {
                    'file_name': file.filename,
                    'source_type':".docx"
                }
            final_list.append(Document(page_content=content, metadata=metadata))
        return final_list
        
    except Exception as e:
       error_message = type(e).__name__
       traceback_message = traceback.format_exc()
       logging.error("Error: %s\nTraceback: %s", error_message, traceback_message)
       return []
